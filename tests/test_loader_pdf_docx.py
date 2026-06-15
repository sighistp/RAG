"""TDD: Tests for PDF and DOCX loading."""
from unittest.mock import patch
import os
from pathlib import Path
from rag.loader import load, load_pdf, load_docx


def _create_test_docx(path: str, text: str = "Hello World"):
    """Create a simple DOCX file for testing."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


# ── load_pdf ──────────────────────────────────────────────────────

@patch("opendataloader_pdf.convert")
def test_load_pdf_extracts_text(mock_convert, tmp_path):
    """load_pdf should convert PDF and return extracted markdown."""
    pdf = tmp_path / "doc.pdf"
    pdf.write_text("dummy", encoding="utf-8")
    out_dir = tmp_path / "out"
    out_dir.mkdir()
    out_file = out_dir / "doc.md"
    out_file.write_text("# Extracted Title\n\nSome content.", encoding="utf-8")

    # Mock convert to write the output file
    def _fake_convert(input_path, output_dir, **kw):
        os.makedirs(output_dir, exist_ok=True)
        Path(output_dir, "doc.md").write_text("# Extracted Title\n\nSome content.", encoding="utf-8")
    mock_convert.side_effect = _fake_convert

    result = load_pdf(str(pdf))
    assert "# Extracted Title" in result
    assert "Some content." in result


@patch("opendataloader_pdf.convert")
def test_load_pdf_called_with_correct_args(mock_convert, tmp_path):
    """load_pdf should call convert with format=markdown."""
    pdf = tmp_path / "doc.pdf"
    pdf.write_text("dummy", encoding="utf-8")

    def _fake_convert(input_path, output_dir, **kw):
        os.makedirs(output_dir, exist_ok=True)
        Path(output_dir, "doc.md").write_text("content", encoding="utf-8")

    mock_convert.side_effect = _fake_convert

    load_pdf(str(pdf))

    args, kwargs = mock_convert.call_args
    assert kwargs.get("format") == "markdown"


@patch("opendataloader_pdf.convert")
def test_load_pdf_missing_file(mock_convert, tmp_path):
    """load_pdf should raise FileNotFoundError for nonexistent PDF."""
    import pytest
    with pytest.raises(FileNotFoundError):
        load_pdf(str(tmp_path / "nonexistent.pdf"))


# ── load_docx ──────────────────────────────────────────────────────

def test_load_docx_extracts_text(tmp_path):
    """load_docx should extract text from DOCX paragraphs."""
    docx = tmp_path / "doc.docx"
    _create_test_docx(str(docx), "Hello World")
    result = load_docx(str(docx))
    assert "Hello World" in result


def test_load_docx_multiple_paragraphs(tmp_path):
    """load_docx should join paragraphs with newlines."""
    docx = tmp_path / "doc.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Para 1")
    doc.add_paragraph("Para 2")
    doc.add_paragraph("Para 3")
    doc.save(str(docx))
    result = load_docx(str(docx))
    assert "Para 1" in result
    assert "Para 2" in result
    assert "Para 3" in result


def test_load_docx_empty(tmp_path):
    """load_docx should return empty string for empty DOCX."""
    docx = tmp_path / "empty.docx"
    from docx import Document
    doc = Document()
    doc.save(str(docx))
    result = load_docx(str(docx))
    assert result == ""


def test_load_docx_missing_file(tmp_path):
    """load_docx should raise FileNotFoundError."""
    import pytest
    with pytest.raises(FileNotFoundError):
        load_docx(str(tmp_path / "nonexistent.docx"))


def test_load_docx_extracts_table_text(tmp_path):
    """load_docx should extract text from tables."""
    from docx import Document
    docx = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "30"
    doc.add_paragraph("After table")
    doc.save(str(docx))
    result = load_docx(str(docx))
    assert "Before table" in result
    assert "After table" in result
    assert "Name" in result
    assert "Age" in result
    assert "Alice" in result
    assert "30" in result


# ── load() dispatch ────────────────────────────────────────────────

@patch("opendataloader_pdf.convert")
def test_load_dispatches_to_pdf(mock_convert, tmp_path):
    """load() should call load_pdf for .pdf files."""
    pdf = tmp_path / "test.pdf"
    pdf.write_text("dummy", encoding="utf-8")

    def _fake_convert(input_path, output_dir, **kw):
        os.makedirs(output_dir, exist_ok=True)
        Path(output_dir, "test.md").write_text("pdf content", encoding="utf-8")

    mock_convert.side_effect = _fake_convert

    result = load(str(pdf))
    assert "pdf content" in result


def test_load_dispatches_to_docx(tmp_path):
    """load() should call load_docx for .docx files."""
    docx = tmp_path / "test.docx"
    _create_test_docx(str(docx), "docx content")
    result = load(str(docx))
    assert "docx content" in result


def test_load_txt_still_works(tmp_path):
    """load() should still work for .txt files."""
    txt = tmp_path / "test.txt"
    txt.write_text("plain text", encoding="utf-8")
    result = load(str(txt))
    assert result == "plain text"


def test_load_md_still_works(tmp_path):
    """load() should still work for .md files."""
    md = tmp_path / "test.md"
    md.write_text("# markdown", encoding="utf-8")
    result = load(str(md))
    assert result == "# markdown"
